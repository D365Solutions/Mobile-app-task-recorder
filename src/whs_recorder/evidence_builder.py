import os
import json
import datetime
from typing import List, Tuple, Optional

import cv2
import numpy as np
from docx import Document
from docx.shared import Inches


def _sharpness(frame) -> float:
    g = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    return cv2.Laplacian(g, cv2.CV_64F).var()


def _edge_density(frame) -> float:
    g = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    g = cv2.GaussianBlur(g, (3, 3), 0)
    e = cv2.Canny(g, 60, 140)
    return float((e > 0).mean())


def _get_frame_at(cap, frame_index: int):
    cap.set(cv2.CAP_PROP_POS_FRAMES, max(int(frame_index), 0))
    ok, frame = cap.read()
    return frame if ok else None


def _best_frame_near(cap, fps: float, t_sec: float, window_sec: float,
                     sharp_min: float, edge_min: float) -> Optional[np.ndarray]:
    center = int(t_sec * fps)
    radius = max(int(window_sec * fps), 1)
    step = max(int(fps // 4), 1)

    best = None
    best_score = -1.0

    for fi in range(max(center - radius, 0), center + radius + 1, step):
        frame = _get_frame_at(cap, fi)
        if frame is None:
            continue

        s = _sharpness(frame)
        ed = _edge_density(frame)
        if s < sharp_min or ed < edge_min:
            continue

        score = s * (1.0 + 10.0 * ed)
        if score > best_score:
            best_score = score
            best = frame

    return best


def _choose_frame(cap, fps: float, t_sec: float) -> Tuple[Optional[np.ndarray], str]:
    # strict
    f = _best_frame_near(cap, fps, t_sec, window_sec=1.8, sharp_min=65, edge_min=0.012)
    if f is not None:
        return f, "strict"
    # relaxed
    f = _best_frame_near(cap, fps, t_sec, window_sec=2.5, sharp_min=35, edge_min=0.008)
    if f is not None:
        return f, "relaxed"
    # exact fallback
    f = _get_frame_at(cap, int(t_sec * fps))
    if f is not None:
        return f, "exact"
    return None, "none"


def build_evidence(
    video: str,
    markers: str,
    out_dir: str,
    title: str = "WHS Mobile – Test Evidence",
    skip_loading: bool = True,
    result_offsets: List[float] = None,
):
    """
    Builds a Word doc from:
      - MP4 screen recording
      - step_markers.json with title/notes/is_loading

    For each marker:
      - capture 1 "action" screenshot around marker time
      - capture 1 "result" screenshot slightly AFTER marker (offsets)
    """
    if result_offsets is None:
        result_offsets = [0.6, 1.2]

    os.makedirs(out_dir, exist_ok=True)

    with open(markers, "r", encoding="utf-8") as f:
        data = json.load(f)
    ms = data.get("markers", [])
    if not ms:
        raise RuntimeError("No markers found.")

    cap = cv2.VideoCapture(video)
    if not cap.isOpened():
        raise RuntimeError(f"Cannot open video: {video}")

    fps = cap.get(cv2.CAP_PROP_FPS) or 10.0

    run_stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    run_out = os.path.join(out_dir, f"evidence_{run_stamp}")
    os.makedirs(run_out, exist_ok=True)

    captures = []  # list of dicts per step

    step_no = 0
    for m in ms:
        if skip_loading and m.get("is_loading") is True:
            continue

        step_no += 1
        t = float(m.get("t", 0.0))
        step_title = (m.get("title") or f"Step {step_no}").strip()
        notes = (m.get("notes") or "").strip()

        # ACTION frame near marker
        action_frame, action_mode = _choose_frame(cap, fps, t)
        action_path = None
        if action_frame is not None:
            action_path = os.path.join(run_out, f"step_{step_no:02d}_action_{action_mode}.jpg")
            cv2.imwrite(action_path, action_frame, [int(cv2.IMWRITE_JPEG_QUALITY), 92])

        # RESULT frame AFTER marker (to catch toast/success message)
        best_result = None
        best_meta = ("none", None)
        for off in result_offsets:
            f, mode = _choose_frame(cap, fps, t + off)
            if f is None:
                continue
            score = _sharpness(f) * (1.0 + 10.0 * _edge_density(f))
            if best_result is None:
                best_result = f
                best_meta = (mode, off)
            else:
                prev_score = _sharpness(best_result) * (1.0 + 10.0 * _edge_density(best_result))
                if score > prev_score:
                    best_result = f
                    best_meta = (mode, off)

        result_path = None
        if best_result is not None:
            mode, off = best_meta
            result_path = os.path.join(run_out, f"step_{step_no:02d}_result_{mode}_plus{off:.1f}s.jpg")
            cv2.imwrite(result_path, best_result, [int(cv2.IMWRITE_JPEG_QUALITY), 92])

        captures.append({
            "no": step_no,
            "t": t,
            "title": step_title,
            "notes": notes,
            "action_img": action_path,
            "result_img": result_path,
        })

    cap.release()

    # Build Word doc
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Video: {os.path.basename(video)}")
    doc.add_paragraph(f"Markers: {os.path.basename(markers)}")
    doc.add_paragraph("Per step: action screenshot + result screenshot (after marker) when available.")

    doc.add_heading("Steps", level=2)

    for c in captures:
        doc.add_heading(f"{c['no']}. {c['title']}", level=3)
        doc.add_paragraph(f"(Marker ~{c['t']:.1f}s)")
        if c["notes"]:
            doc.add_paragraph(c["notes"])

        if c["action_img"]:
            doc.add_paragraph("Action:")
            doc.add_picture(c["action_img"], width=Inches(6.3))

        if c["result_img"]:
            doc.add_paragraph("Result:")
            doc.add_picture(c["result_img"], width=Inches(6.3))

    out_doc = os.path.join(run_out, "WHS_Test_Evidence.docx")
    doc.save(out_doc)
    print("Created:", out_doc)
    return out_doc
